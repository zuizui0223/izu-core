#!/usr/bin/env python3
"""Audit PLOS supporting files from the four Canary–Balearic communities.

The audit separates community identity from data scope. The legacy Word files
store tables as embedded vector pictures, so plain text extraction is attempted
first and a LibreOffice -> PDF -> pdftotext route is used when the document body
contains only picture placeholders.
"""
from __future__ import annotations

import argparse
import json
import math
import re
import shutil
import subprocess
import tempfile
from collections import Counter
from pathlib import Path
from typing import Any, Iterable, Sequence


COMMUNITY_CODES = ("SB", "CM", "CB", "LC")
PAIR_TERMS = ("flower visitor", "flower-visitor", "pollinator", "visitor")
WEIGHT_TERMS = (
    "fvr",
    "flower visitation rate",
    "visit frequency",
    "interaction weight",
)
DERIVED_TERMS = (
    "linkage level",
    "selectiveness",
    "specificity",
    "functional richness",
    "rank abundance",
    "evenness of abundances",
)
DERIVED_TABLE_CORE = (
    "specificity",
    "zone",
    "sp.cod",
    "month",
    "functional",
    "richness",
    "rank",
    "abundance",
    "evenness",
)
SPECIES_CODE_PATTERN = re.compile(
    r"(?<![A-Za-z0-9])([a-z]{3}\.[a-z0-9]{2,3})(?![A-Za-z0-9])"
)


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def normalize_space(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def _decode_output(payload: bytes) -> str:
    for encoding in ("utf-8", "cp1252", "latin-1"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


def _run_text_command(command: Sequence[str]) -> str:
    result = subprocess.run(command, capture_output=True, check=False)
    if result.returncode != 0 or not result.stdout:
        raise RuntimeError(
            f"command failed: {' '.join(command)}; returncode={result.returncode}; "
            f"stderr={result.stderr[:300]!r}"
        )
    return _decode_output(result.stdout)


def needs_render_fallback(text: str) -> bool:
    """Detect antiword output that contains only embedded-picture placeholders."""
    nonempty = [normalize_space(line) for line in text.splitlines() if line.strip()]
    picture_markers = sum("[pic]" in line.casefold() for line in nonempty)
    substantive = [line for line in nonempty if "[pic]" not in line.casefold()]
    return picture_markers > 0 and len(substantive) <= max(4, picture_markers + 1)


def extract_doc_with_antiword(path: Path) -> str:
    errors: list[str] = []
    for command in (["antiword", "-t", str(path)], ["antiword", str(path)]):
        try:
            return _run_text_command(command)
        except Exception as error:
            errors.append(repr(error))
    raise RuntimeError("; ".join(errors))


def extract_doc_via_pdf(path: Path) -> str:
    office = next(
        (command for command in ("libreoffice", "soffice") if shutil.which(command)),
        None,
    )
    if office is None:
        raise RuntimeError(
            "LibreOffice/soffice is required for image-only legacy Word tables"
        )
    if shutil.which("pdftotext") is None:
        raise RuntimeError("pdftotext is required for rendered legacy Word tables")
    with tempfile.TemporaryDirectory(prefix="cb_plos_doc_") as temporary:
        output_dir = Path(temporary)
        result = subprocess.run(
            [
                office,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(path),
            ],
            capture_output=True,
            check=False,
        )
        pdf_path = output_dir / f"{path.stem}.pdf"
        if result.returncode != 0 or not pdf_path.exists():
            raise RuntimeError(
                f"LibreOffice conversion failed for {path}; "
                f"returncode={result.returncode}; stdout={result.stdout[:300]!r}; "
                f"stderr={result.stderr[:300]!r}"
            )
        return _run_text_command(["pdftotext", "-layout", str(pdf_path), "-"])


def extract_doc(path: Path) -> tuple[str, str]:
    antiword_text = extract_doc_with_antiword(path)
    if not needs_render_fallback(antiword_text):
        return antiword_text, "antiword"
    return extract_doc_via_pdf(path), "libreoffice_pdf_pdftotext"


def extract_docx(path: Path) -> tuple[str, str]:
    from docx import Document

    document = Document(path)
    lines: list[str] = []
    lines.extend(
        paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()
    )
    for table_index, table in enumerate(document.tables, start=1):
        lines.append(f"[TABLE {table_index}]")
        for row in table.rows:
            lines.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(lines), "python_docx"


def extract_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.casefold()
    if suffix == ".doc":
        return extract_doc(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix in {".txt", ".csv", ".tsv"}:
        return (
            path.read_text(encoding="utf-8-sig", errors="replace"),
            "plain_text",
        )
    raise ValueError(f"unsupported supporting-file type: {path}")


def logical_id_from_name(name: str) -> str | None:
    match = re.search(r"(?:^|[._-])(s00[1-9])(?:[._-]|$)", name.casefold())
    return match.group(1) if match else None


def species_codes(text: str) -> list[str]:
    return sorted(set(SPECIES_CODE_PATTERN.findall(text.casefold())))


def is_number(value: str) -> bool:
    cleaned = value.strip().replace(",", ".")
    if not cleaned:
        return False
    try:
        return math.isfinite(float(cleaned))
    except ValueError:
        return False


def split_fields(line: str) -> list[str]:
    if "\t" in line:
        fields = [normalize_space(value) for value in line.split("\t")]
    else:
        fields = [
            normalize_space(value) for value in re.split(r"\s{2,}", line.strip())
        ]
    return [field for field in fields if field]


def table_like_lines(lines: Sequence[str]) -> list[dict[str, object]]:
    output: list[dict[str, object]] = []
    for index, line in enumerate(lines, start=1):
        fields = split_fields(line)
        if len(fields) < 3:
            continue
        numeric_after_first = sum(is_number(value) for value in fields[1:])
        output.append(
            {
                "line": index,
                "n_fields": len(fields),
                "numeric_fraction_after_first": (
                    numeric_after_first / max(1, len(fields) - 1)
                ),
                "preview": fields[:12],
            }
        )
    return output


def numeric_matrix_blocks(
    rows: Sequence[dict[str, object]],
) -> list[dict[str, object]]:
    candidates = [
        row
        for row in rows
        if int(row["n_fields"]) >= 4
        and float(row["numeric_fraction_after_first"]) >= 0.7
        and row.get("preview")
        and not is_number(str(row["preview"][0]))
    ]
    blocks: list[list[dict[str, object]]] = []
    current: list[dict[str, object]] = []
    for row in candidates:
        if current and int(row["line"]) != int(current[-1]["line"]) + 1:
            if len(current) >= 3:
                blocks.append(current)
            current = []
        current.append(row)
    if len(current) >= 3:
        blocks.append(current)
    return [
        {
            "start_line": int(block[0]["line"]),
            "end_line": int(block[-1]["line"]),
            "n_rows": len(block),
            "typical_field_count": Counter(
                int(row["n_fields"]) for row in block
            ).most_common(1)[0][0],
            "preview": [row["preview"] for row in block[:3]],
        }
        for block in blocks
    ]


def count_codes(text: str) -> dict[str, int]:
    return {
        code: len(
            re.findall(
                rf"(?<![A-Za-z0-9]){re.escape(code)}(?![A-Za-z0-9])",
                text,
            )
        )
        for code in COMMUNITY_CODES
    }


def candidate_lines(
    lines: Sequence[str], terms: Iterable[str]
) -> list[dict[str, object]]:
    lowered_terms = tuple(term.casefold() for term in terms)
    output = []
    for index, line in enumerate(lines, start=1):
        lowered = line.casefold()
        matched = [term for term in lowered_terms if term in lowered]
        if matched:
            output.append(
                {
                    "line": index,
                    "matched_terms": matched,
                    "text": normalize_space(line)[:500],
                }
            )
    return output[:40]


def has_pairwise_header(lines: Sequence[str]) -> bool:
    for line in lines:
        fields = split_fields(line)
        if len(fields) < 3:
            continue
        lowered = " | ".join(fields).casefold()
        has_plant = "plant" in lowered
        has_visitor = any(term in lowered for term in PAIR_TERMS)
        has_weight = any(term in lowered for term in WEIGHT_TERMS)
        if has_plant and has_visitor and has_weight:
            return True
    return False


def is_derived_partner_summary(text: str) -> bool:
    lowered = normalize_space(text).casefold()
    return all(term in lowered for term in DERIVED_TABLE_CORE)


def overlap_summary(query: set[str], reference: set[str]) -> dict[str, object]:
    shared = sorted(query & reference)
    return {
        "n_query_codes": len(query),
        "n_reference_codes": len(reference),
        "n_shared_codes": len(shared),
        "query_coverage": len(shared) / len(query) if query else None,
        "shared_codes": shared,
        "query_only_codes": sorted(query - reference),
        "reference_only_codes": sorted(reference - query),
    }


def audit_file(
    path: Path,
    *,
    output_dir: Path,
    declared_scope: str | None,
) -> dict[str, object]:
    text, method = extract_text(path)
    text_dir = output_dir / "text"
    text_dir.mkdir(parents=True, exist_ok=True)
    text_path = text_dir / f"{path.name}.txt"
    text_path.write_text(text, encoding="utf-8")
    lines = [line.rstrip() for line in text.splitlines() if line.strip()]
    rows = table_like_lines(lines)
    derived = is_derived_partner_summary(text)
    blocks = [] if derived else numeric_matrix_blocks(rows)
    pair_headers = has_pairwise_header(lines)
    code_counts = count_codes(text)
    lowered = text.casefold()
    codes = species_codes(text)
    return {
        "logical_id": logical_id_from_name(path.name),
        "filename": path.name,
        "declared_scope": declared_scope,
        "text_extraction_method": method,
        "text_path": str(text_path.relative_to(output_dir)),
        "n_nonempty_text_lines": len(lines),
        "community_code_counts": code_counts,
        "community_codes_present": [
            code for code, count in code_counts.items() if count > 0
        ],
        "species_codes": codes,
        "n_species_codes": len(codes),
        "contains_fvr_term": any(term in lowered for term in WEIGHT_TERMS),
        "contains_derived_partner_terms": [
            term for term in DERIVED_TERMS if term in lowered
        ],
        "derived_species_partner_summary": derived,
        "pairwise_interaction_header_candidate": pair_headers,
        "n_table_like_lines": len(rows),
        "numeric_matrix_blocks": blocks,
        "n_numeric_matrix_blocks": len(blocks),
        "content_term_lines": candidate_lines(
            lines,
            (*PAIR_TERMS, *WEIGHT_TERMS, *DERIVED_TERMS, "zone", "month"),
        ),
        "opening_preview": [normalize_space(line)[:500] for line in lines[:20]],
    }


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--config",
        type=Path,
        default=Path("config/canary_balearic_plos_source.json"),
    )
    parser.add_argument(
        "--input-dir",
        type=Path,
        default=Path("artifacts/canary_balearic_plos"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("artifacts/canary_balearic_plos/content_audit.json"),
    )
    args = parser.parse_args()

    config = load_json(args.config)
    source_inventory = load_json(args.input_dir / "source_inventory.json")
    expected_scope = {
        str(row["logical_id"]): str(row.get("declared_scope") or "")
        for row in config.get("expected_files", [])
    }
    audits: list[dict[str, object]] = []
    errors: list[dict[str, str]] = []
    file_paths = sorted(
        path
        for path in (args.input_dir / "files").rglob("*")
        if path.is_file()
        and path.suffix.casefold() in {".doc", ".docx", ".txt", ".csv", ".tsv"}
    )
    for path in file_paths:
        logical_id = logical_id_from_name(path.name)
        try:
            audits.append(
                audit_file(
                    path,
                    output_dir=args.input_dir,
                    declared_scope=expected_scope.get(logical_id or ""),
                )
            )
        except Exception as error:
            errors.append({"file": str(path), "error": repr(error)})

    by_id = {
        str(row["logical_id"]): row for row in audits if row.get("logical_id")
    }
    aggregate_code_counts = {
        code: sum(
            int(row["community_code_counts"].get(code, 0)) for row in audits
        )
        for code in COMMUNITY_CODES
    }
    present_codes = [
        code for code, count in aggregate_code_counts.items() if count > 0
    ]
    raw_pair_candidates = [
        row for row in audits if bool(row["pairwise_interaction_header_candidate"])
    ]
    numeric_matrix_candidates = [
        row for row in audits if int(row["n_numeric_matrix_blocks"]) > 0
    ]
    same_four_communities_supported = set(present_codes) == set(COMMUNITY_CODES)
    derived_summary_files = [
        row for row in audits if bool(row["derived_species_partner_summary"])
    ]

    plant_reference = set(by_id.get("s001", {}).get("species_codes", []))
    visitor_reference = set(by_id.get("s002", {}).get("species_codes", []))
    s3_codes = set(by_id.get("s003", {}).get("species_codes", []))
    s4_codes = set(by_id.get("s004", {}).get("species_codes", []))
    code_domain_crosswalk = {
        "s003_vs_s001_plant_codes": overlap_summary(s3_codes, plant_reference),
        "s003_vs_s002_visitor_codes": overlap_summary(
            s3_codes, visitor_reference
        ),
        "s004_vs_s001_plant_codes": overlap_summary(s4_codes, plant_reference),
        "s004_vs_s002_visitor_codes": overlap_summary(
            s4_codes, visitor_reference
        ),
    }
    s3_visitor_coverage = code_domain_crosswalk[
        "s003_vs_s002_visitor_codes"
    ]["query_coverage"]
    s3_plant_coverage = code_domain_crosswalk[
        "s003_vs_s001_plant_codes"
    ]["query_coverage"]
    s4_plant_coverage = code_domain_crosswalk[
        "s004_vs_s001_plant_codes"
    ]["query_coverage"]
    s4_visitor_coverage = code_domain_crosswalk[
        "s004_vs_s002_visitor_codes"
    ]["query_coverage"]
    caption_content_role_reversal = bool(
        s3_visitor_coverage is not None
        and s3_plant_coverage is not None
        and s4_plant_coverage is not None
        and s4_visitor_coverage is not None
        and s3_visitor_coverage > s3_plant_coverage
        and s4_plant_coverage > s4_visitor_coverage
    )

    if (
        same_four_communities_supported
        and not raw_pair_candidates
        and not numeric_matrix_candidates
    ):
        if derived_summary_files:
            scope_status = (
                "same_four_communities_supported_"
                "derived_species_partner_summaries_only"
            )
        else:
            scope_status = "same_four_communities_supported_no_raw_network_candidate"
    elif same_four_communities_supported and (
        raw_pair_candidates or numeric_matrix_candidates
    ):
        scope_status = (
            "same_four_communities_with_raw_network_candidate_manual_mapping_required"
        )
    elif audits:
        scope_status = "supporting_files_recovered_community_identity_incomplete"
    else:
        scope_status = "supporting_file_content_not_readable"

    summary = {
        "schema_version": "1.1",
        "status": scope_status,
        "source_id": config["source_id"],
        "article_doi": config["article_doi"],
        "pmcid": config["pmcid"],
        "target_source_id": config.get("target_source_id"),
        "target_article_doi": config.get("target_article_doi"),
        "source_inventory_status": source_inventory.get("status"),
        "n_supporting_files_audited": len(audits),
        "n_text_extraction_errors": len(errors),
        "aggregate_community_code_counts": aggregate_code_counts,
        "community_codes_present": present_codes,
        "same_four_communities_supported": same_four_communities_supported,
        "n_derived_species_partner_summary_files": len(derived_summary_files),
        "n_pairwise_interaction_header_candidates": len(raw_pair_candidates),
        "n_numeric_matrix_candidate_files": len(numeric_matrix_candidates),
        "code_domain_crosswalk": code_domain_crosswalk,
        "caption_content_role_mismatch_detected": (
            caption_content_role_reversal
        ),
        "resolved_content_roles": (
            {
                "s001": "selected_plant_species_metadata",
                "s002": "selected_flower_visitor_species_metadata",
                "s003": "flower_visitor_by_month_derived_partner_traits",
                "s004": "plant_by_month_derived_partner_traits",
            }
            if caption_content_role_reversal
            else {}
        ),
        "caption_content_role_reading": (
            "The article labels S3 as the plant dataset and S4 as the "
            "flower-visitor dataset, but the source-code crosswalk shows S3 "
            "codes align with the S2 visitor list and S4 codes align with the "
            "S1 plant list."
            if caption_content_role_reversal
            else "No source-code role reversal was established."
        ),
        "full_network_source_admitted": False,
        "effect_registry_eligible": False,
        "audits": audits,
        "errors": errors,
        "reading": (
            "All four named community codes are present. S1 and S2 list "
            "selected plant and visitor species, while S3 and S4 contain "
            "species-by-month derived partner summaries (L or d-prime, "
            "functional richness, rank abundance and evenness), not "
            "plant-by-visitor edges or full quantitative matrices."
            if same_four_communities_supported and derived_summary_files
            else (
                "The supporting files establish only the scope reported by the "
                "audit; no missing interaction matrix is reconstructed."
            )
        ),
        "next_gate": (
            "Retain this package as a source-locked same-community derivative "
            "dataset and continue searching for the original 2014 full "
            "matrices or source-native community metrics. Do not enter these "
            "derived rows into the cross-system effect registry."
            if scope_status
            == "same_four_communities_supported_derived_species_partner_summaries_only"
            else (
                "If a raw interaction candidate is present, write an explicit "
                "source-specific column and community mapping; otherwise "
                "continue searching for the original matrices."
            )
        ),
        "claim_boundary": config["claim_boundary"],
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(
        json.dumps(summary, indent=2, ensure_ascii=False, default=str) + "\n",
        encoding="utf-8",
    )
    print(f"supporting files audited: {len(audits)}")
    print(f"community codes: {present_codes}")
    print(f"derived summary files: {len(derived_summary_files)}")
    print(f"raw pair candidates: {len(raw_pair_candidates)}")
    print(f"numeric matrix candidates: {len(numeric_matrix_candidates)}")
    print(f"caption/content role reversal: {caption_content_role_reversal}")
    print(f"status: {scope_status}")


if __name__ == "__main__":
    main()
